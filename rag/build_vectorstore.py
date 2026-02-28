"""
AuditMind — RAG Vector Store Builder
Embeds regulatory documents into ChromaDB for semantic retrieval.
Agents call query_regulations() to get cited regulatory passages.
"""

import os
import re
import json
import pandas as pd
from pathlib import Path
from tqdm import tqdm

import chromadb
from chromadb.utils import embedding_functions

# ── Paths ────────────────────────────────────────────────────────────────────
ROOT       = Path(__file__).parent.parent
RAW_DOCS   = ROOT / "data/raw/regulatory_docs"
PROCESSED  = ROOT / "data/processed"
VECTOR_DIR = ROOT / "rag/vectorstore"
VECTOR_DIR.mkdir(parents=True, exist_ok=True)


# ════════════════════════════════════════════════════════════════════════════
# 1.  IMPROVED CHUNKER  (fixes the 1-chunk-per-DOCX problem from Step 3)
# ════════════════════════════════════════════════════════════════════════════

def chunk_by_paragraph(text: str, source: str,
                        min_words: int = 40,
                        max_words: int = 200,
                        overlap_words: int = 30) -> list[dict]:
    """
    Chunks text using paragraph boundaries as primary split points,
    then merges short paragraphs and splits long ones.
    This produces much better RAG chunks than naive word-window splitting.
    """
    chunks = []
    chunk_id = 0

    # Clean text
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)

    # Split on paragraph boundaries (double newline, or sentence-ending periods
    # followed by capital letters — catches regulatory numbered sections)
    raw_paragraphs = re.split(r'\n{2,}|(?<=[.!?])\s{2,}(?=[A-Z])', text)
    raw_paragraphs = [p.strip() for p in raw_paragraphs if len(p.strip()) > 20]

    # Merge short paragraphs together until we hit min_words
    merged = []
    buffer = ""
    for para in raw_paragraphs:
        buffer = (buffer + " " + para).strip()
        if len(buffer.split()) >= min_words:
            merged.append(buffer)
            buffer = ""
    if buffer:  # Don't lose the last bit
        if merged:
            merged[-1] += " " + buffer
        else:
            merged.append(buffer)

    # Split merged blocks that exceed max_words (with overlap)
    for block in merged:
        words = block.split()
        if len(words) <= max_words:
            chunks.append({
                "chunk_id":  f"{source}_{chunk_id:04d}",
                "source":    source,
                "text":      block,
                "word_count": len(words)
            })
            chunk_id += 1
        else:
            # Sliding window with overlap
            i = 0
            while i < len(words):
                window = words[i:i + max_words]
                chunk_text = " ".join(window)
                if len(window) >= min_words // 2:
                    chunks.append({
                        "chunk_id":  f"{source}_{chunk_id:04d}",
                        "source":    source,
                        "text":      chunk_text,
                        "word_count": len(window)
                    })
                    chunk_id += 1
                i += max_words - overlap_words

    return chunks


def extract_and_chunk_all_docs() -> list[dict]:
    """
    Re-extracts and re-chunks all regulatory documents.
    Fixes the 1-chunk-per-DOCX problem by using paragraph-aware chunking.
    """
    try:
        import fitz
        from docx import Document
    except ImportError:
        os.system("pip install PyMuPDF python-docx -q")
        import fitz
        from docx import Document

    all_chunks = []

    # ── PDFs ─────────────────────────────────────────────────────────────────
    for pdf_path in sorted(RAW_DOCS.glob("*.pdf")):
        print(f"  Chunking PDF: {pdf_path.name}")
        try:
            doc = fitz.open(str(pdf_path))
            # Extract with page breaks preserved (better paragraph detection)
            full_text = ""
            for page in doc:
                full_text += page.get_text() + "\n\n"
            chunks = chunk_by_paragraph(full_text, source=pdf_path.stem)
            all_chunks.extend(chunks)
            print(f"    → {len(chunks)} chunks (was: limited before)")
        except Exception as e:
            print(f"    ✗ Error: {e}")

    # ── DOCX files ───────────────────────────────────────────────────────────
    for docx_path in sorted(RAW_DOCS.glob("*.docx")):
        print(f"  Chunking DOCX: {docx_path.name}")
        try:
            doc = Document(str(docx_path))
            # Extract each paragraph separately, then rechunk
            paragraphs = [
                p.text.strip()
                for p in doc.paragraphs
                if len(p.text.strip()) > 15
            ]
            # Also extract from tables (FFIEC docs store content in tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if len(text) > 15:
                            paragraphs.append(text)

            full_text = "\n\n".join(paragraphs)
            chunks = chunk_by_paragraph(full_text, source=docx_path.stem)
            all_chunks.extend(chunks)
            print(f"    → {len(chunks)} chunks (was: 1 before)")
        except Exception as e:
            print(f"    ✗ Error: {e}")
    
    # Save for inspection
    df = pd.DataFrame(all_chunks)
    df.to_csv(PROCESSED / "regulatory_chunks_v2.csv", index=False)
    print(f"\n  Total chunks: {len(all_chunks)}")
    print(f"  Avg words/chunk: {df['word_count'].mean():.0f}")
    print(f"  Per source:\n{df.groupby('source')['chunk_id'].count().to_string()}")

    return all_chunks


# ════════════════════════════════════════════════════════════════════════════
# 2.  BUILD CHROMADB VECTOR STORE
# ════════════════════════════════════════════════════════════════════════════

def build_vectorstore(chunks: list[dict]):
    """
    Embeds all chunks and stores in ChromaDB.
    Uses sentence-transformers (free, runs locally, no API key needed).
    Model: all-MiniLM-L6-v2 — fast, good quality, 384-dim embeddings.
    """
    print(f"\n{'='*60}")
    print("  Building ChromaDB Vector Store")
    print("='*60}")

    # Use sentence-transformers — runs locally, no API cost
    # First run downloads the model (~80MB, cached after that)
    print("  Loading embedding model (all-MiniLM-L6-v2)...")
    print("  First run downloads ~80MB model. This is cached after.")
    ef = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name="all-MiniLM-L6-v2"
    )

    # Create persistent ChromaDB client
    client = chromadb.PersistentClient(path=str(VECTOR_DIR))

    # Delete existing collection if rebuilding
    try:
        client.delete_collection("regulations")
        print("  Deleted existing collection (rebuilding)")
    except Exception:
        pass

    collection = client.create_collection(
        name="regulations",
        embedding_function=ef,
        metadata={"description": "FATF + FFIEC regulatory corpus for AuditMind"}
    )

    # Embed in batches (ChromaDB handles large sets better in batches)
    batch_size = 50
    total = len(chunks)
    print(f"  Embedding {total} chunks in batches of {batch_size}...")

    for i in tqdm(range(0, total, batch_size), desc="  Embedding"):
        batch = chunks[i:i + batch_size]
        collection.add(
            ids=[c["chunk_id"] for c in batch],
            documents=[c["text"] for c in batch],
            metadatas=[{"source": c["source"], "word_count": c["word_count"]} for c in batch]
        )

    count = collection.count()
    print(f"\n  ✓ Vector store built: {count} embeddings")
    print(f"  ✓ Saved to: {VECTOR_DIR}")
    return collection


# ════════════════════════════════════════════════════════════════════════════
# 3.  RETRIEVAL FUNCTION  (this is what agents will call)
# ════════════════════════════════════════════════════════════════════════════

def query_regulations(query: str,
                      n_results: int = 3,
                      source_filter: str = None) -> list[dict]:
    """
    Main retrieval function called by agents.

    Args:
        query: Natural language question or topic
        n_results: How many passages to return (default 3)
        source_filter: Optional — restrict to one source
                       e.g. 'FATF_40_Recommendations'

    Returns:
        List of dicts with 'text', 'source', 'relevance_score'

    Example:
        results = query_regulations("structuring threshold $10000")
        for r in results:
            print(r['source'], r['text'][:200])
    """
    ef = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name="all-MiniLM-L6-v2"
    )
    client = chromadb.PersistentClient(path=str(VECTOR_DIR))
    collection = client.get_collection("regulations", embedding_function=ef)

    # Optional source filter
    where = {"source": source_filter} if source_filter else None

    results = collection.query(
        query_texts=[query],
        n_results=n_results,
        where=where,
        include=["documents", "metadatas", "distances"]
    )

    # Format results for agent consumption
    formatted = []
    for doc, meta, dist in zip(
        results["documents"][0],
        results["metadatas"][0],
        results["distances"][0]
    ):
        # Convert distance to similarity score (lower distance = higher similarity)
        similarity = round(1 - dist, 4)
        formatted.append({
            "text":             doc,
            "source":           meta["source"],
            "relevance_score":  similarity,
            "citation":         format_citation(meta["source"])
        })

    return formatted


def format_citation(source: str) -> str:
    """Maps source filename to a human-readable regulatory citation"""
    citations = {
        "FATF_40_Recommendations":        "FATF 40 Recommendations (2012, updated 2024)",
        "FATF_Risk_Based_Approach":       "FATF Guidance: Risk-Based Approach for Financial Institutions",
        "FFIEC_Customer_Due_Diligence":   "FFIEC BSA/AML Manual: Customer Due Diligence (2018)",
        "FFIEC_Suspicious_Activity_Reporting": "FFIEC BSA/AML Manual: Suspicious Activity Reporting (2014)",
        "FFIEC_Currency_Transaction_Reporting":"FFIEC BSA/AML Manual: Currency Transaction Reporting (2021)",
        "FFIEC_BSA_AML_Risk_Assessment":  "FFIEC BSA/AML Manual: Risk Assessment (2020)",
        "FFIEC_Internal_Controls":        "FFIEC BSA/AML Manual: Internal Controls (2020)",
    }
    return citations.get(source, source.replace("_", " "))


# ════════════════════════════════════════════════════════════════════════════
# 4.  TEST QUERIES  (validates retrieval quality after building)
# ════════════════════════════════════════════════════════════════════════════

def test_retrieval():
    """
    Runs 5 test queries covering the main use cases.
    Prints results so you can visually verify quality.
    """
    print(f"\n{'='*60}")
    print("  Testing Retrieval Quality")
    print(f"{'='*60}")

    test_queries = [
        ("Structuring transactions below $10,000 threshold",
         "Should retrieve FATF structuring / threshold guidance"),

        ("Customer due diligence requirements for high risk clients",
         "Should retrieve FFIEC CDD or FATF Rec. 10"),

        ("Suspicious activity reporting filing requirements",
         "Should retrieve FFIEC SAR procedures"),

        ("Money laundering placement layering integration",
         "Should retrieve FATF ML typologies"),

        ("AML compliance program internal controls",
         "Should retrieve FFIEC internal controls"),
    ]

    for query, expected in test_queries:
        print(f"\n  Query: '{query}'")
        print(f"  Expected: {expected}")
        results = query_regulations(query, n_results=2)
        for i, r in enumerate(results, 1):
            print(f"\n  Result {i} (score: {r['relevance_score']:.3f})")
            print(f"  Source: {r['citation']}")
            print(f"  Text preview: {r['text'][:180]}...")

    print(f"\n{'='*60}")
    print("  ✓ Retrieval test complete")
    print(f"{'='*60}")


# ════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("\n" + "█"*60)
    print("  AuditMind — RAG Vector Store Builder")
    print("█"*60)

    print("\n[1/3] Extracting and chunking documents...")
    chunks = extract_and_chunk_all_docs()

    print("\n[2/3] Building ChromaDB vector store...")
    collection = build_vectorstore(chunks)

    print("\n[3/3] Testing retrieval quality...")
    test_retrieval()

    print("\n" + "█"*60)
    print("  ✓ RAG vector store ready")
    print(f"  Location: {VECTOR_DIR}")
    print("  Import with: from rag.build_vectorstore import query_regulations")
    print("█"*60 + "\n")